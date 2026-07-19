import subprocess
import zipfile
from pathlib import Path

from docx import Document

from manuscript_placeholders.collector import collect


def test_markdown_sections_and_categories(tmp_path: Path) -> None:
    source = tmp_path / "story.md"
    source.write_text("# Chapter One\nMara entered [NAME THIS TAVERN].\nTODO: improve this exchange\n# Chapter Two\n[RESEARCH inheritance law]\n", encoding="utf-8")
    items = collect(source)
    assert [item.category for item in items] == ["Missing names", "Revision notes", "Research"]
    assert items[0].section == "Chapter One"
    assert items[2].section == "Chapter Two"


def test_docx_headings_are_preserved(tmp_path: Path) -> None:
    source = tmp_path / "story.docx"
    document = Document()
    document.add_heading("Chapter Seven", level=1)
    document.add_paragraph("The room was cold. <insert argument here>")
    document.save(source)
    items = collect(source)
    assert len(items) == 1
    assert items[0].section == "Chapter Seven"
    assert items[0].category == "Insertions"


def test_overlapping_bracket_patterns_are_not_duplicated(tmp_path: Path) -> None:
    source = tmp_path / "story.txt"
    source.write_text("[RESEARCH THIS CUSTOM]\n", encoding="utf-8")
    assert len(collect(source)) == 1


def test_custom_pattern(tmp_path: Path) -> None:
    source = tmp_path / "story.txt"
    source.write_text("FLAGME later\n", encoding="utf-8")
    items = collect(source, (r"FLAGME",))
    assert items[0].category == "Custom"


def test_epub_spine_and_headings_are_preserved(tmp_path: Path) -> None:
    source = tmp_path / "story.epub"
    container = '<?xml version="1.0"?><container xmlns="urn:oasis:names:tc:opendocument:xmlns:container"><rootfiles><rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/></rootfiles></container>'
    package = '<?xml version="1.0"?><package xmlns="http://www.idpf.org/2007/opf"><manifest><item id="chapter" href="chapter.xhtml" media-type="application/xhtml+xml"/></manifest><spine><itemref idref="chapter"/></spine></package>'
    chapter = '<html xmlns="http://www.w3.org/1999/xhtml"><body><h1>Chapter Nine</h1><p>She stopped. [RESEARCH local custom]</p></body></html>'
    with zipfile.ZipFile(source, "w") as archive:
        archive.writestr("mimetype", "application/epub+zip")
        archive.writestr("META-INF/container.xml", container)
        archive.writestr("OEBPS/content.opf", package)
        archive.writestr("OEBPS/chapter.xhtml", chapter)
    items = collect(source)
    assert len(items) == 1
    assert items[0].section == "Chapter Nine"
    assert items[0].category == "Research"


def test_legacy_doc_uses_local_reader(monkeypatch, tmp_path: Path) -> None:
    source = tmp_path / "story.doc"
    source.write_bytes(b"legacy fixture")
    monkeypatch.setattr("manuscript_placeholders.collector.shutil.which", lambda tool: "antiword" if tool == "antiword" else None)
    monkeypatch.setattr(
        "manuscript_placeholders.collector.subprocess.run",
        lambda *args, **kwargs: subprocess.CompletedProcess(args[0], 0, stdout=b"TODO: revise the ending\n", stderr=b""),
    )
    items = collect(source)
    assert len(items) == 1
    assert items[0].category == "Revision notes"


def test_legacy_doc_explains_reader_requirement(monkeypatch, tmp_path: Path) -> None:
    source = tmp_path / "story.doc"
    source.write_bytes(b"legacy fixture")
    monkeypatch.setattr("manuscript_placeholders.collector.shutil.which", lambda tool: None)
    try:
        collect(source)
    except ValueError as exc:
        assert "antiword, catdoc, or LibreOffice" in str(exc)
    else:
        raise AssertionError("Expected a local reader requirement")
